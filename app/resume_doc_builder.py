"""
resume_doc_builder.py
----------------------
Small helper used only by the Gmail send flow (Step 5).

Problem: the "attach resume" feature on the send modal attaches whatever
file the user uploaded in resume_file — but if they instead pasted their
resume as plain text (pasted_resume_text), there's no underlying file to
attach at all.

This module closes that gap: it turns pasted resume text into a plain
.docx, purely as a formatting exercise. It does NOT call the AI, does NOT
rewrite or improve anything, and does NOT invent content — it renders
exactly what the user pasted, one paragraph per line, so the attachment
matches what they typed to the letter. If they want a nicer-looking resume
document, they should upload one instead; this is a fallback so a send
never goes out with no resume attached at all when text is available.
"""

import io

from docx import Document

GENERATED_RESUME_FILENAME = "Resume.docx"
GENERATED_RESUME_MIME = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)


class ResumeDocBuildError(Exception):
    pass


def build_resume_docx(resume_text: str) -> bytes:
    """
    Renders pasted resume text as a simple .docx (one paragraph per line,
    blank lines preserved for readability). Raises ResumeDocBuildError if
    there's no usable text to render.
    """
    if not resume_text or not resume_text.strip():
        raise ResumeDocBuildError("No pasted resume text to build a document from.")

    doc = Document()
    for line in resume_text.splitlines():
        # add_paragraph("") keeps blank lines as visual spacing rather than
        # collapsing them, since resumes rely on that spacing for structure.
        doc.add_paragraph(line.rstrip())

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
