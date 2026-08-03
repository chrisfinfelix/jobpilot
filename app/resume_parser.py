"""
resume_parser.py
-----------------
Extracts raw text from an uploaded resume file, AND (Step 4) runs
deterministic, non-AI ATS compatibility checks on it.

Deliberately keeping both in this file rather than a separate module:
ATS structural issues (tables, missing sections, no contact info, scanned
images) are direct byproducts of *how the file parses* — which is exactly
what this file already knows better than anything else in the app.
Splitting it out would mean parsing the resume twice for no reason.

The fuzzier ATS concern — "does this resume use the right keywords for
THIS job" — needs actual judgment, not a fixed rule, so that part stays
in groq_client.py's prompt instead of here.
"""

import io
import re
from docx import Document
import pdfplumber

# Section headers most ATS software is trained to look for. Not exhaustive —
# just common enough that a resume missing ALL of them is worth flagging.
EXPECTED_SECTIONS = {
    "experience": ["experience", "work history", "employment"],
    "education": ["education", "academic"],
    "skills": ["skills", "technical skills", "competencies"],
}

EMAIL_PATTERN = re.compile(r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}")
PHONE_PATTERN = re.compile(r"(\+?\d[\d\-\.\s\(\)]{8,}\d)")

MIN_WORD_COUNT = 120   # below this, a resume is likely missing real content
MAX_WORD_COUNT = 1200  # above this, it's likely too long for a 1-2 page resume


class UnsupportedFileType(Exception):
    """Raised when the uploaded file extension isn't one we support."""
    pass


class EmptyResumeError(Exception):
    """Raised when we couldn't extract any usable text from the file."""
    pass


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract text from a PDF file's raw bytes."""
    text_chunks = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_chunks.append(page_text)
    return "\n".join(text_chunks)


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from a .docx file's raw bytes."""
    doc = Document(io.BytesIO(file_bytes))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

    # Also pull text out of any tables (some resumes use table layouts)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text)

    return "\n".join(paragraphs)


def parse_resume(filename: str | None, file_bytes: bytes | None, pasted_text: str | None) -> str:
    """
    Main entry point. Accepts EITHER an uploaded file OR pasted text and
    returns clean extracted text.

    Priority: if pasted_text is provided, use it directly (no parsing needed).
    Otherwise, look at the filename extension to decide how to parse the file.
    """
    if pasted_text and pasted_text.strip():
        text = pasted_text.strip()
    elif filename and file_bytes:
        ext = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""
        if ext == "pdf":
            text = extract_text_from_pdf(file_bytes)
        elif ext == "docx":
            text = extract_text_from_docx(file_bytes)
        elif ext == "txt":
            text = file_bytes.decode("utf-8", errors="ignore")
        else:
            raise UnsupportedFileType(
                f"'.{ext}' is not supported. Please upload a PDF, DOCX, or paste text."
            )
    else:
        raise EmptyResumeError("No resume file or pasted text was provided.")

    text = text.strip()
    if len(text) < 30:
        # A real resume is always longer than this. This usually means
        # the PDF was a scanned image with no selectable text layer.
        raise EmptyResumeError(
            "Couldn't extract readable text from this resume. "
            "If it's a scanned/image-based PDF, try pasting the text instead."
        )

    return text


def _uses_table_layout(file_bytes: bytes) -> bool:
    """DOCX-only: true if a meaningful chunk of content lives inside tables."""
    try:
        doc = Document(io.BytesIO(file_bytes))
    except Exception:
        return False

    table_char_count = sum(
        len(cell.text) for table in doc.tables for row in table.rows for cell in row.cells
    )
    paragraph_char_count = sum(len(p.text) for p in doc.paragraphs)
    total = table_char_count + paragraph_char_count
    if total == 0:
        return False
    return (table_char_count / total) > 0.25  # more than a quarter of content in tables


def _has_embedded_images(file_bytes: bytes) -> bool:
    """PDF-only: true if the PDF contains images (possible scan, photo, icon graphics)."""
    try:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            return any(len(page.images) > 0 for page in pdf.pages)
    except Exception:
        return False


def check_ats_structural_issues(
    filename: str | None,
    file_bytes: bytes | None,
    extracted_text: str,
) -> list[str]:
    """
    Step 4: deterministic, non-AI ATS compatibility checks. These are rule-based
    (not model judgment) because they're objectively true or false — whether a
    resume has a detectable "Experience" section isn't something worth asking
    an LLM to guess at when we can just check.

    Returns a list of plain-English warning strings. Empty list = no structural
    issues detected.
    """
    warnings: list[str] = []
    ext = filename.lower().rsplit(".", 1)[-1] if filename and "." in filename else None

    # --- File format ---
    if ext == "txt":
        warnings.append(
            "Plain .txt files lose all formatting — most ATS systems expect a PDF or DOCX upload."
        )

    # --- Section headers ---
    lowered = extracted_text.lower()
    missing_sections = [
        label for label, keywords in EXPECTED_SECTIONS.items()
        if not any(kw in lowered for kw in keywords)
    ]
    if len(missing_sections) == len(EXPECTED_SECTIONS):
        warnings.append(
            "No standard section headers detected (e.g. Experience, Education, Skills) — "
            "ATS software often relies on these to categorize your content."
        )
    elif missing_sections:
        warnings.append(
            f"Missing a clearly labeled '{missing_sections[0].title()}' section header — "
            "consider adding one so ATS software can find it."
        )

    # --- Contact info ---
    if not EMAIL_PATTERN.search(extracted_text):
        warnings.append(
            "No email address detected in the extracted text — if it's inside a header/footer "
            "graphic, some ATS systems won't read it. Put contact info in the main body."
        )
    if not PHONE_PATTERN.search(extracted_text):
        warnings.append("No phone number detected in the extracted text.")

    # --- Length ---
    word_count = len(extracted_text.split())
    if word_count < MIN_WORD_COUNT:
        warnings.append(
            f"Resume text is quite short (~{word_count} words) — it may be missing content "
            "an ATS or recruiter would expect to see."
        )
    elif word_count > MAX_WORD_COUNT:
        warnings.append(
            f"Resume text is long (~{word_count} words) — consider trimming to fit 1-2 pages."
        )

    # --- File-specific structural checks ---
    if ext == "docx" and file_bytes and _uses_table_layout(file_bytes):
        warnings.append(
            "This DOCX appears to use a table-based layout — some ATS systems fail to read "
            "text inside tables correctly, or read it out of order."
        )
    if ext == "pdf" and file_bytes and _has_embedded_images(file_bytes):
        warnings.append(
            "This PDF contains images or graphics (e.g. a photo, icons, or a graphical layout) — "
            "these are typically invisible to ATS text scanning."
        )

    return warnings